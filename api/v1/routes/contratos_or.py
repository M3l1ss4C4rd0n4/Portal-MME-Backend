"""
Endpoint contratos OR — Portal Dirección EE
GET /v1/contratos-or/dashboard → contratos_or.seguimiento_avance_documental (Power BI logic)
GET /v1/contratos-or/actas/catalogo → Lista actas PDF desde SharePoint
GET /v1/contratos-or/actas/descargar → Descarga acta PDF por itemId
"""

import logging
from io import BytesIO
from fastapi import APIRouter, Depends, HTTPException, Request
from fastapi.responses import JSONResponse, Response
from slowapi import Limiter
from slowapi.util import get_remote_address

from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet

from api.dependencies import get_api_key
from infrastructure.database.connection import PostgreSQLConnectionManager

from .reports import _sp_open_folder, _graph_list_children_recent, _sp_download_item

logger = logging.getLogger(__name__)
router = APIRouter()
limiter = Limiter(key_func=get_remote_address)

_cm = PostgreSQLConnectionManager()


@router.get("/dashboard", summary="Dashboard seguimiento contratos OR")
@limiter.limit("60/minute")
async def get_contratos_or_dashboard(request: Request, api_key: str = Depends(get_api_key)):
    try:
        with _cm.get_connection() as conn:
            with conn.cursor() as cur:
                # KPIs por proyecto
                cur.execute("""
                    WITH pagados AS (
                        SELECT nombre_proyecto_id, nro_desembolso
                        FROM contratos_or.seguimiento_avance_documental
                        WHERE necesaria_para_desembolso = 'Sí'
                        GROUP BY nombre_proyecto_id, nro_desembolso
                        HAVING COUNT(*) = COUNT(CASE WHEN estado = 'Completo' THEN 1 END)
                    )
                    SELECT
                        s.nombre_proyecto_id,
                        s.ejecutor,
                        s.departamento,
                        s.municipio,
                        ROUND(AVG(s.avance) * 100, 1)                                 AS avance_general,
                        ROUND(COALESCE(SUM(DISTINCT
                            CASE WHEN p.nro_desembolso IS NOT NULL AND s.desembolso IS NOT NULL
                                 THEN s.desembolso ELSE 0 END
                        ), 0) * 100, 1)                                               AS avance_financiero,
                        COUNT(DISTINCT CASE WHEN s.nro_desembolso IS NOT NULL
                              THEN s.nro_desembolso END)                               AS total_desembolsos,
                        COUNT(DISTINCT p.nro_desembolso)                               AS pagos_realizados
                    FROM contratos_or.seguimiento_avance_documental s
                    LEFT JOIN pagados p
                        ON s.nombre_proyecto_id = p.nombre_proyecto_id
                       AND s.nro_desembolso     = p.nro_desembolso
                       AND s.necesaria_para_desembolso = 'Sí'
                    GROUP BY s.nombre_proyecto_id, s.ejecutor, s.departamento, s.municipio
                    ORDER BY s.nombre_proyecto_id
                """)
                cols_p = [d[0] for d in cur.description]
                proyectos_rows = [dict(zip(cols_p, r)) for r in cur.fetchall()]

                # KPIs globales desde la hoja Resumen del Excel
                cur.execute("""
                    SELECT
                        MAX(CASE WHEN unnamed_0 = 'Numero de Contratos' THEN unnamed_1 END) AS n_contratos,
                        MAX(CASE WHEN unnamed_0 = 'Avance Financiero' THEN unnamed_1 END)   AS avance_financiero,
                        MAX(CASE WHEN unnamed_0 = 'Avance General' THEN unnamed_1 END)      AS avance_general,
                        MAX(CASE WHEN unnamed_0 = 'Pagos Realizados' THEN unnamed_1 END)    AS pagos_realizados,
                        MAX(CASE WHEN unnamed_0 = 'Total Pagos' THEN unnamed_1 END)         AS total_pagos,
                        MAX(CASE WHEN unnamed_0 = '% Pagos' THEN unnamed_1 END)             AS pct_pagos
                    FROM contratos_or.resumen
                """)
                g = cur.fetchone()

                cur.execute("SELECT MAX(fecha_carga) FROM contratos_or.resumen")
                ts_or = cur.fetchone()[0]

                # Avance Físico desde la hoja Seguimiento_Avance_Fisico
                cur.execute("""
                    SELECT ROUND(AVG(avance), 1) AS avance_fisico
                    FROM contratos_or.seguimiento_avance_fisico
                    WHERE avance IS NOT NULL
                """)
                af = cur.fetchone()[0]

                # Per-project avance_fisico
                cur.execute("""
                    SELECT nombre_proyecto_id, ROUND(AVG(avance), 1) AS avance_fisico
                    FROM contratos_or.seguimiento_avance_fisico
                    WHERE avance IS NOT NULL
                      AND nombre_proyecto_id IS NOT NULL
                      AND nombre_proyecto_id != 'Promedio Avance'
                    GROUP BY nombre_proyecto_id
                """)
                af_per_project = {row[0]: float(row[1]) if row[1] is not None else 0.0
                                 for row in cur.fetchall()}

                # Progreso por desembolso
                cur.execute("""
                    WITH pagados AS (
                        SELECT nombre_proyecto_id, nro_desembolso
                        FROM contratos_or.seguimiento_avance_documental
                        WHERE necesaria_para_desembolso = 'Sí'
                        GROUP BY nombre_proyecto_id, nro_desembolso
                        HAVING COUNT(*) = COUNT(CASE WHEN estado = 'Completo' THEN 1 END)
                    )
                    SELECT
                        s.nro_desembolso::int                                          AS numero,
                        COUNT(CASE WHEN s.necesaria_para_desembolso = 'Sí' THEN 1 END) AS act_necesarias,
                        COUNT(CASE WHEN s.necesaria_para_desembolso = 'Sí'
                                    AND s.estado = 'Completo' THEN 1 END)              AS act_completas,
                        ROUND(
                            COUNT(CASE WHEN s.necesaria_para_desembolso = 'Sí'
                                        AND s.estado = 'Completo' THEN 1 END)::numeric /
                            NULLIF(COUNT(CASE WHEN s.necesaria_para_desembolso = 'Sí' THEN 1 END), 0) * 100
                        , 1)                                                           AS pct_completado,
                        COUNT(DISTINCT p.nombre_proyecto_id)                          AS proyectos_pagados
                    FROM contratos_or.seguimiento_avance_documental s
                    LEFT JOIN pagados p
                        ON s.nombre_proyecto_id = p.nombre_proyecto_id
                       AND s.nro_desembolso     = p.nro_desembolso
                    WHERE s.nro_desembolso IS NOT NULL
                    GROUP BY s.nro_desembolso
                    ORDER BY s.nro_desembolso
                """)
                cols_d = [d[0] for d in cur.description]
                desembolsos_rows = [dict(zip(cols_d, r)) for r in cur.fetchall()]

        def _f(v):
            return float(v) if v is not None else None

        n_contratos       = int(g[0] or 0)
        avance_financiero = _f(g[1]) * 100 if g[1] is not None else 0.0  # 0.5->50%
        avance_general    = _f(g[2]) * 100 if g[2] is not None else 0.0  # 0.25->25%
        pagos_real        = int(g[3] or 0)
        pagos_pos         = int(g[4] or 0)
        avance_fisico     = _f(af) or 0.0  # ya en % desde la DB

        return JSONResponse({
            "ultima_actualizacion": ts_or.strftime("%d/%m/%Y, %H:%M") if ts_or else None,
            "nContratos":        n_contratos,
            "avanceGeneral":     avance_general,
            "avanceFinanciero":  avance_financiero,
            "pagosRealizados":   pagos_real,
            "pagosPosibles":     pagos_pos,
            "pctPagosRealizados": round(pagos_real / pagos_pos * 100, 1) if pagos_pos else 0.0,
            "avanceFisico":      avance_fisico,
            "desembolsos": [
                {
                    "numero":            int(d["numero"]),
                    "actNecesarias":     int(d["act_necesarias"]),
                    "actCompletas":      int(d["act_completas"]),
                    "pctCompletado":     _f(d["pct_completado"]),
                    "proyectosPagados":  int(d["proyectos_pagados"]),
                }
                for d in desembolsos_rows
            ],
            "proyectos": [
                {
                    "nombre":           p["nombre_proyecto_id"],
                    "ejecutor":         p["ejecutor"],
                    "departamento":     p["departamento"],
                    "municipio":        p["municipio"],
                    "avanceGeneral":    _f(p["avance_general"]),
                    "avanceFisico":     af_per_project.get(p["nombre_proyecto_id"], 0.0),
                    "avanceFinanciero": _f(p["avance_financiero"]),
                    "totalDesembolsos": int(p["total_desembolsos"]),
                    "pagosRealizados":  int(p["pagos_realizados"]),
                }
                for p in proyectos_rows
            ],
        })
    except Exception as e:
        logger.error("[contratos-or] %s", e)
        raise HTTPException(status_code=500, detail="Error al obtener contratos OR")


# ─── Helper: convertir .docx a PDF ────────────────────────────────────────
def _convert_docx_to_pdf_bytes(docx_bytes: bytes) -> bytes:
    """Convierte un archivo .docx en bytes a PDF usando python-docx + reportlab."""
    doc = Document(BytesIO(docx_bytes))
    pdf_buffer = BytesIO()
    doc_template = SimpleDocTemplate(pdf_buffer, pagesize=letter)
    styles = getSampleStyleSheet()
    flowables: list = []

    # Párrafos
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            style_name = "Heading1" if para.style and para.style.name == "Heading 1" else "Normal"
            s = styles[style_name] if style_name in styles else styles["Normal"]
            try:
                flowables.append(Paragraph(text, s))
                flowables.append(Spacer(1, 6))
            except Exception:
                flowables.append(Paragraph(text.replace("&", "&amp;"), styles["Normal"]))
                flowables.append(Spacer(1, 6))

    # Tablas
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            row_text = " | ".join(cells)
            if row_text.strip():
                flowables.append(Paragraph(row_text, styles["Normal"]))
                flowables.append(Spacer(1, 3))
        flowables.append(Spacer(1, 8))

    doc_template.build(flowables)
    return pdf_buffer.getvalue()


# ─── SharePoint: Actas de seguimiento ELECTROCAQUETA ────────────────────────

SHARE_URL_ACTAS = (
    "https://minenergiacol.sharepoint.com/:f:/r/sites/"
    "DireccindeEnergaElctrica-DEE_Supervision/Shared%20Documents/"
    "DEE_Supervision/Direccion_Energia/1%20-%20Fondos%20de%20Inversi%C3%B3n/"
    "10.%20COMUNIDADES%20ENERG%C3%89TICAS/ELECTROCAQUETA%20ACTAS"
    "?csf=1&web=1&e=jVfXf9"
)


@router.get("/actas/catalogo", summary="Catálogo de actas de seguimiento")
@limiter.limit("30/minute")
async def get_actas_catalogo(request: Request, api_key: str = Depends(get_api_key)):
    try:
        headers, drive_id, root_id = _sp_open_folder(SHARE_URL_ACTAS)
        children = _graph_list_children_recent(headers, drive_id, root_id, top=50)

        documentos = []
        for item in children:
            if not item.get("file"):
                continue
            name_lower = item["name"].lower()
            if not (name_lower.endswith(".pdf") or name_lower.endswith(".docx")):
                continue
            tipo = "pdf" if name_lower.endswith(".pdf") else "docx"
            documentos.append({
                "id": item["id"],
                "name": item["name"],
                "lastModified": item.get("lastModifiedDateTime"),
                "tipo": tipo,
            })

        return {"actas": documentos, "total": len(documentos)}
    except Exception as e:
        logger.error("[contratos-or/actas/catalogo] %s", e)
        raise HTTPException(status_code=502, detail="Error al obtener actas de SharePoint")


@router.get("/actas/descargar", summary="Descargar acta PDF")
@limiter.limit("30/minute")
async def descargar_acta(
    request: Request,
    item_id: str,
    api_key: str = Depends(get_api_key),
):
    try:
        headers, drive_id, root_id = _sp_open_folder(SHARE_URL_ACTAS)

        # Obtener nombre del archivo
        children = _graph_list_children_recent(headers, drive_id, root_id, top=50)
        filename = "acta.pdf"
        for child in children:
            if child["id"] == item_id:
                filename = child.get("name", "acta.pdf")
                break

        file_bytes = _sp_download_item(headers, drive_id, {"id": item_id, "name": filename})

        # Si es .docx, convertir a PDF antes de enviar
        if filename.lower().endswith(".docx"):
            try:
                file_bytes = _convert_docx_to_pdf_bytes(file_bytes)
                filename = filename.rsplit(".", 1)[0] + ".pdf"
            except Exception as conv_err:
                logger.error("[contratos-or/actas/descargar] Error convirtiendo docx a PDF: %s", conv_err)
                # Si falla la conversión, intentamos enviar el docx original
                return Response(
                    content=file_bytes,
                    media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    headers={
                        "Content-Disposition": f'attachment; filename="{filename}"',
                        "Cache-Control": "no-store, no-cache, must-revalidate",
                    },
                )

        return Response(
            content=file_bytes,
            media_type="application/pdf",
            headers={
                "Content-Disposition": f'attachment; filename="{filename}"',
                "Cache-Control": "no-store, no-cache, must-revalidate",
            },
        )
    except Exception as e:
        logger.error("[contratos-or/actas/descargar] %s", e)
        raise HTTPException(status_code=502, detail="Error al descargar acta")
